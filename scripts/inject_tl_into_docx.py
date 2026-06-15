"""Injecte les sections Transfer Learning (8.7 -> 8.12) dans le rapport Word, juste
avant 'Conclusions scientifiques et metiers', pour produire UN rapport final complet
et editable. Numerotation auto (styles Heading 2) -> Word renumerote tout seul.
Source des chiffres : notebook Rattrapage (1).ipynb. Modele retenu : ResNet50.
"""
import os
import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "RAPPORT FINAL VF (2).docx")
OUT = os.path.join(ROOT, "RAPPORT FINAL VF.docx")
FIG = os.path.join(ROOT, "reports", "figures")
BLUE = RGBColor(0x1F, 0x6F, 0xB2)

doc = docx.Document(SRC)

# Point d'insertion : le titre "Conclusions scientifiques et metiers"
target = None
for p in doc.paragraphs:
    if "Conclusions scientifiques" in p.text:
        target = p._p
        break
if target is None:
    raise SystemExit("Titre 'Conclusions scientifiques' introuvable.")


def _before(el):
    target.addprevious(el)


def h2(text):
    p = doc.add_paragraph(text, style="Heading 2")
    _before(p._p)


def sub(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = BLUE
    _before(p._p)


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _before(p._p)


def bullets(items):
    for it in items:
        try:
            p = doc.add_paragraph(it, style="List Bullet")
        except KeyError:
            p = doc.add_paragraph("• " + it)
        _before(p._p)


def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
    _before(t._tbl)
    # paragraphe vide apres le tableau
    _before(doc.add_paragraph()._p)


def image(name, width_cm=15.0, caption=None):
    path = os.path.join(FIG, name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    _before(p._p)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        _before(c._p)


# ====================== 8.7 METHODOLOGIE ======================
h2("Transfer Learning – Méthodologie et justification des choix")
body("Cette section détaille et justifie chaque choix de paramètre : découpage des données, "
     "prétraitement, augmentation, architecture de la tête, gestion du déséquilibre et protocole "
     "de fine-tuning.")

sub("a)  Pourquoi le Transfer Learning")
body("Le CNN entraîné de zéro (section précédente) généralise mal (accuracy de validation "
     "≈ 50 %, macro F1-score 0,20) : avec un jeu limité, il prédit massivement la classe majoritaire "
     "Normal. Le Transfer Learning réutilise des réseaux pré-entraînés sur ImageNet "
     "(1,2 million d’images) : leurs premières couches détectent déjà des motifs génériques "
     "(contours, textures) qu’il suffit d’adapter à la radiographie. C’est l’état de l’art pour la "
     "classification d’images médicales avec peu de données.")

sub("b)  Reproductibilité (graine aléatoire 42)")
body("La graine est fixée à 42 pour NumPy, TensorFlow et les générateurs d’images : même "
     "découpage, même ordre de mélange et mêmes tirages d’augmentation à chaque exécution. C’est "
     "indispensable pour comparer les modèles à armes égales et pour reproduire et défendre exactement "
     "les chiffres présentés.")

sub("c)  Pourquoi un découpage stratifié 70 / 15 / 15")
body("Pourquoi stratifier ? Le jeu est fortement déséquilibré (Normal 48 %, Viral Pneumonia 6 %). "
     "Un découpage purement aléatoire risquerait de sous- ou sur-représenter une classe rare dans le "
     "test, rendant la mesure instable. La stratification (paramètre stratify sur le label) garantit que "
     "chaque ensemble (train, validation, test) conserve les mêmes proportions de classes que le jeu "
     "complet.")
body("Pourquoi 70 / 15 / 15 ? 70 % pour l’entraînement (assez de données pour fine-tuner) ; 15 % de "
     "validation, pour le réglage et les callbacks (EarlyStopping, ReduceLR) sans jamais intervenir dans "
     "la mesure finale ; 15 % de test, jamais vus, pour l’évaluation finale. Le découpage se fait en "
     "deux temps (70 % / 30 %, puis ce 30 % coupé en deux), avec random_state = 42 (déterministe). La "
     "séparation porte sur les images : une image ne peut appartenir qu’à un seul ensemble → pas de "
     "fuite de données.")
table([
    ["Classe", "Train", "Validation", "Test"],
    ["COVID", "2 531", "543", "542"],
    ["Lung Opacity", "4 208", "902", "902"],
    ["Normal", "7 134", "1 529", "1 529"],
    ["Viral Pneumonia", "942", "201", "202"],
    ["Total", "14 815", "3 175", "3 175"],
])

sub("d)  Taille d’image 224×224 et conversion RGB")
body("Les images sont redimensionnées en 224×224 car ResNet50 et EfficientNetB0 ont été "
     "pré-entraînés à cette résolution ; la conserver permet de réutiliser pleinement les poids "
     "ImageNet. Les radiographies étant en niveaux de gris, elles sont chargées en 3 canaux (RGB) par "
     "réplication du canal : les poids ImageNet attendent 3 canaux. On ne fabrique pas de couleur, on "
     "respecte le format d’entrée.")

sub("e)  Normalisation par preprocess_input dédié")
body("Chaque backbone applique sa propre normalisation (resnet50.preprocess_input ≠ "
     "efficientnet.preprocess_input), celle attendue par ses poids pré-entraînés, et non un simple "
     "/255 générique : sinon les statistiques d’entrée ne correspondraient pas à celles vues pendant "
     "l’entraînement ImageNet, dégradant les performances.")

sub("f)  Augmentation de données (train uniquement) – valeurs justifiées")
body("L’augmentation n’est appliquée qu’au train, pour accroître sa diversité et limiter le "
     "sur-apprentissage ; validation et test restent intacts. Chaque valeur reste réaliste en radiologie :")
bullets([
    "rotation ±10° : un patient n’est jamais parfaitement aligné, mais une grande rotation serait irréaliste sur un thorax.",
    "décalages ±8 % (horizontal et vertical) : tolérance au cadrage.",
    "zoom ±10 % : variation de distance / cadrage.",
    "retournement horizontal : un thorax reste plausible en miroir gauche/droite.",
    "pas de retournement vertical : une radio retournée (cœur en haut) n’a aucun sens anatomique.",
    "fill_mode = nearest : les pixels créés par rotation/décalage sont comblés par le voisin le plus proche, évitant des bandes noires artificielles.",
])

sub("g)  Architecture de la tête de classification")
body("Sur le backbone gelé, on ajoute : GlobalAveragePooling2D → BatchNormalization → Dropout(0,4) "
     "→ Dense(128, ReLU) → Dropout(0,3) → Dense(4, Softmax). Le GlobalAveragePooling résume les "
     "cartes de caractéristiques en un vecteur compact (moins de paramètres qu’un Flatten, donc moins de "
     "sur-apprentissage). Les deux Dropout et la BatchNormalization régularisent et stabilisent "
     "l’apprentissage. La sortie Softmax à 4 neurones donne une probabilité par classe.")

sub("h)  Taille de lot (batch = 32)")
body("Compromis classique : assez grand pour des gradients stables et une bonne utilisation du matériel, "
     "assez petit pour tenir en mémoire. 32 est une valeur standard et robuste.")

sub("i)  Gestion du déséquilibre – poids de classe « balanced »")
body("Des poids de classe sont calculés avec compute_class_weight(\"balanced\") sur le train uniquement "
     "(anti-fuite) : chaque classe reçoit un poids inversement proportionnel à sa fréquence. Le modèle "
     "est donc davantage pénalisé lorsqu’il se trompe sur une classe rare, ce qui l’empêche de « tout "
     "prédire Normal ».")
table([
    ["Classe", "Poids"],
    ["COVID", "1,463"], ["Lung Opacity", "0,880"],
    ["Normal", "0,519"], ["Viral Pneumonia", "3,932"],
])

sub("j)  Fine-tuning en deux phases – pourquoi ce protocole")
bullets([
    "Phase 1 (tête seule, 10 époques, lr = 1e-4) : le backbone est gelé (≈ 267 000 paramètres entraînables sur 23,8 millions pour ResNet50). On laisse la nouvelle tête se stabiliser sans détruire les poids ImageNet par des gradients initiaux trop grands.",
    "Phase 2 (fine-tuning, 15 époques, lr = 1e-5) : on dégèle les 30 dernières couches du backbone (les plus spécifiques), les premières couches restant gelées. On recompile avec un taux d’apprentissage 10× plus faible (1e-5) pour ajuster finement ces couches au domaine médical sans « casser » les poids pré-entraînés.",
])

sub("k)  Callbacks")
bullets([
    "EarlyStopping (patience 5, restauration des meilleurs poids) : arrête l’entraînement si la validation ne progresse plus → évite le sur-apprentissage et conserve le meilleur état.",
    "ModelCheckpoint : sauvegarde automatiquement le meilleur modèle (sur val_loss).",
    "ReduceLROnPlateau (facteur 0,3, patience 2) : réduit le taux d’apprentissage quand la validation stagne, pour affiner la convergence.",
])

sub("l)  Métrique d’évaluation")
body("Vu le déséquilibre, l’accuracy est trompeuse (un classifieur « tout Normal » atteindrait 48 %). "
     "La métrique de sélection est donc le macro F1-score (moyenne non pondérée des F1 par classe : "
     "chaque classe pèse autant), complétée par le rappel par classe (notamment COVID) comme garde-fou "
     "clinique.")

# ====================== 8.8 RESNET50 ======================
h2("Transfer Learning – ResNet50 (modèle retenu)")
body("Évaluation sur le test figé (3 175 images). ResNet50 fine-tuné obtient les meilleures "
     "performances de tous les modèles testés : c’est le modèle retenu.")
sub("Performance globale")
bullets(["Accuracy : 0,933", "Macro F1-score : 0,936", "Précision macro : 0,941",
         "Rappel macro : 0,932", "Weighted F1-score : 0,933"])
sub("Classification report")
table([
    ["Classe", "Précision", "Rappel", "F1-score", "Support"],
    ["COVID", "0,965", "0,958", "0,961", "542"],
    ["Lung Opacity", "0,943", "0,873", "0,906", "902"],
    ["Normal", "0,917", "0,960", "0,938", "1 529"],
    ["Viral Pneumonia", "0,940", "0,936", "0,938", "202"],
    ["Macro avg", "0,941", "0,932", "0,936", "3 175"],
])
body("Observations. ResNet50 fine-tuné atteint un macro F1-score de 0,936, le meilleur de l’étude, "
     "avec un excellent équilibre sur les quatre classes : rappel COVID 0,958 (garde-fou clinique), et la "
     "classe difficile Lung Opacity est la mieux traitée de tous les modèles (rappel 0,873). C’est "
     "pourquoi il est retenu comme modèle final.")

# ====================== 8.9 EFFICIENTNET ======================
h2("Transfer Learning – EfficientNetB0")
body("Évaluation sur le test figé (3 175 images).")
sub("Performance globale")
bullets(["Accuracy : 0,893", "Macro F1-score : 0,896", "Précision macro : 0,881",
         "Rappel macro : 0,913", "Weighted F1-score : 0,893"])
sub("Classification report")
table([
    ["Classe", "Précision", "Rappel", "F1-score", "Support"],
    ["COVID", "0,836", "0,924", "0,878", "542"],
    ["Lung Opacity", "0,860", "0,900", "0,880", "902"],
    ["Normal", "0,939", "0,869", "0,902", "1 529"],
    ["Viral Pneumonia", "0,890", "0,960", "0,924", "202"],
    ["Macro avg", "0,881", "0,913", "0,896", "3 175"],
])
body("Observations. EfficientNetB0 fine-tuné atteint lui aussi un très bon niveau (macro F1 0,896), "
     "avec un bon rappel macro (0,913) et un rappel COVID de 0,924. Il reste légèrement en deçà de "
     "ResNet50 sur le F1-macro global, ce qui le place en second modèle de l’étude.")

# ====================== 8.10 COMPARAISON ======================
h2("Comparaison des modèles")
body("Les deux modèles de Transfer Learning sont évalués sur le test figé (3 175 images) ; le CNN "
     "seul et l’hybride CNN + Gradient Boosting ont été évalués sur le découpage 80/20 et sont "
     "rappelés à titre indicatif.")
table([
    ["Modèle", "Macro F1", "Accuracy", "Rappel COVID", "Test"],
    ["ResNet50 (retenu)", "0,936", "0,933", "0,958", "figé · 3 175"],
    ["EfficientNetB0", "0,896", "0,893", "0,924", "figé · 3 175"],
    ["Hybride CNN + GB", "0,83", "0,83", "0,80", "80/20"],
    ["CNN seul", "0,20", "0,50", "0,00", "80/20"],
])
body("Lecture. Le Transfer Learning améliore très nettement la classification par rapport au CNN "
     "entraîné de zéro (macro F1 0,20, qui ne détecte aucun COVID) et à l’approche hybride (0,83). "
     "ResNet50 fine-tuné est le meilleur modèle (macro F1 0,936) et constitue le modèle retenu, suivi de "
     "près par EfficientNetB0 (0,896). Cela valide l’apport des architectures pré-entraînées sur "
     "ImageNet pour ce problème.")

# ====================== 8.11 INTERPRETABILITE ======================
h2("Interprétabilité étendue – Grad-CAM (Transfer Learning) + SHAP")
body("L’interprétabilité par Grad-CAM est ici étendue aux deux modèles de Transfer Learning et "
     "complétée par SHAP.")
sub("a)  Grad-CAM sur ResNet50 et EfficientNetB0")
body("Grad-CAM est appliqué aux deux modèles, sur la dernière couche convolutive : un exemple "
     "correctement classé par classe (image originale | carte de chaleur | superposition) et des cas mal "
     "classés, pour observer où le modèle regarde lorsqu’il se trompe.")
image("gradcam_resnet50_COVID.png", 15.0, "Grad-CAM – ResNet50, cas COVID correctement classé.")
image("gradcam_efficientnetb0_COVID.png", 15.0, "Grad-CAM – EfficientNetB0, cas COVID correctement classé.")
sub("b)  SHAP")
body("shap.GradientExplainer (robuste sur les modèles fonctionnels Keras récents) est appliqué à un "
     "modèle de Transfer Learning : fond d’environ 24 images d’entraînement, explication de plusieurs "
     "images de test. Cette approche complète Grad-CAM en attribuant à chaque pixel une contribution "
     "(positive ou négative) à la prédiction.")
image("shap_efficientnetb0.png", 15.0,
      "Valeurs SHAP par classe. À gauche la radio originale ; pixels rouges = poussent vers la classe, "
      "bleus = l’en éloignent.")
sub("c)  Lecture et garde-fou")
body("Une activation localisée sur les champs pulmonaires conforte la plausibilité clinique de la "
     "prédiction. À l’inverse, une activation concentrée sur les bords, les marqueurs ou les "
     "annotations signalerait un biais d’apprentissage de raccourci (shortcut learning) : le modèle "
     "exploiterait des artefacts corrélés à la source de l’image plutôt que la pathologie. Les classes "
     "provenant de bases différentes, ce risque est réel et l’interprétabilité sert précisément à "
     "l’auditer.")

# ====================== 8.12 CONCLUSION MODELISATION ======================
h2("Conclusion de la modélisation")
body("La progression Baseline ML → CNN → Gradient Boosting → Transfer Learning montre un gain net et "
     "cohérent. Le modèle retenu est ResNet50 fine-tuné (macro F1 0,936, accuracy 0,933, rappel COVID "
     "0,958), qui obtient les meilleures performances de l’étude ; EfficientNetB0 (0,896) le suit de "
     "près. Le découpage stratifié anti-fuite, les poids de classe, le fine-tuning en deux phases et "
     "l’audit Grad-CAM/SHAP garantissent une évaluation honnête et défendable.")

doc.save(OUT)
print("OK ->", OUT)
print("Paragraphes:", len(docx.Document(OUT).paragraphs))
